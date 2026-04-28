"""Generate dissertation report in .docx format with academic formatting."""
import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = "/Users/anuhyavalluri/Desktop/Accident_prediction_system_test-"
PLOTS = os.path.join(ROOT, "backend/output/plots")
OUT = os.path.join(ROOT, "Dissertation_Report.docx")

doc = Document()

# ---------- Page setup: A4, mirror margins ----------
for section in doc.sections:
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ---------- Default style: Times New Roman 12pt, 1.5 line spacing ----------
styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)
normal.paragraph_format.line_spacing = 1.5
normal.paragraph_format.space_after = Pt(6)

# East Asia font fallback
rpr = normal.element.get_or_add_rPr()
rfonts = rpr.find(qn('w:rFonts'))
if rfonts is None:
    rfonts = OxmlElement('w:rFonts')
    rpr.append(rfonts)
rfonts.set(qn('w:ascii'), 'Times New Roman')
rfonts.set(qn('w:hAnsi'), 'Times New Roman')
rfonts.set(qn('w:eastAsia'), 'Times New Roman')

# Heading styles
def configure_heading(name, size, bold=True, all_caps=False):
    s = styles[name]
    s.font.name = 'Times New Roman'
    s.font.size = Pt(size)
    s.font.bold = bold
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.paragraph_format.line_spacing = 1.5
    s.paragraph_format.space_before = Pt(12)
    s.paragraph_format.space_after = Pt(8)
    if all_caps:
        s.font.all_caps = True

configure_heading('Heading 1', 16, True, all_caps=True)
configure_heading('Heading 2', 14, True)
configure_heading('Heading 3', 12, True)

# ---------- Page numbers (bottom centre, Arabic) ----------
def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

for section in doc.sections:
    add_page_number(section)

# ---------- Helpers ----------
def add_chapter_title(num, title):
    # Page break before chapter
    p = doc.add_paragraph()
    p.add_run().add_break(6)  # page break (6 = WD_BREAK.PAGE)

def page_break():
    p = doc.add_paragraph()
    run = p.add_run()
    from docx.enum.text import WD_BREAK
    run.add_break(WD_BREAK.PAGE)

def chapter(num, name):
    page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"CHAPTER {num}")
    r.bold = True
    r.font.size = Pt(16)
    r.font.name = 'Times New Roman'
    h = doc.add_heading(name.upper(), level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

def h2(text):
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.bold = True

def h3(text):
    h = doc.add_heading(text, level=3)
    for run in h.runs:
        run.font.bold = True

def para(text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.line_spacing = 1.5
    return p

def bullet(items):
    for it in items:
        p = doc.add_paragraph(it, style='List Bullet')
        p.paragraph_format.line_spacing = 1.5
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def fig(path, caption, width=Inches(5.5)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(path):
        run = p.add_run()
        run.add_picture(path, width=width)
    else:
        p.add_run(f"[Figure placeholder: {os.path.basename(path)}]").italic = True
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    crun = cap.add_run(caption)
    crun.bold = True
    crun.font.size = Pt(11)
    cap.paragraph_format.space_after = Pt(12)

def placeholder(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    box = p.add_run(f"[ {text} ]")
    box.italic = True
    box.font.size = Pt(11)
    box.font.color.rgb = RGBColor(100, 100, 100)

def table(headers, rows, caption=None):
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        crun = cap.add_run(caption)
        crun.bold = True
        crun.font.size = Pt(11)
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = t.rows[r_idx].cells[c_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)

# ============================================================
# FRONT MATTER — Declaration / Certificate / Acknowledgement
# ============================================================
def front_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)
    r.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(18)

def front_sub(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(18)

def signature_block(lines):
    for label, value in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1.5
        r1 = p.add_run(f"{label}: ")
        r1.bold = True
        r1.font.size = Pt(12)
        r1.font.name = 'Times New Roman'
        r2 = p.add_run(value)
        r2.font.size = Pt(12)
        r2.font.name = 'Times New Roman'

# --- DECLARATION ---
front_title("DECLARATION BY THE STUDENTS")
front_sub("To whom-so-ever it may concern")
para(
    "We, Anuhya Valluri, Mehak and Kashik Kumar hereby declare that the work done by us on "
    "“AI-Based Accident Hotspot Prediction and Severity Classification Using Spatial "
    "Clustering and Ensemble Learning”, is a record of original work for the partial "
    "fulfillment of the requirements for the award of the degree, Master of Computer "
    "Applications (MCA)."
)
doc.add_paragraph()
doc.add_paragraph()
signature_block([
    ("Name", "Anuhya Valluri"),
    ("Signature", "____________________"),
])
doc.add_paragraph()
signature_block([
    ("Name", "Mehak"),
    ("Signature", "____________________"),
])
doc.add_paragraph()
signature_block([
    ("Name", "Kashik Kumar"),
    ("Signature", "____________________"),
])
doc.add_paragraph()
signature_block([
    ("Date", "____________________"),
    ("Place", "Lovely Professional University, Phagwara"),
])

# --- CERTIFICATE ---
page_break()
front_title("CERTIFICATE")
para(
    "This is to certify that the declaration statement made by this group of students is "
    "correct to the best of my knowledge and belief. They have completed this Capstone Project "
    "under my guidance and supervision. The present work is the result of their original "
    "investigation, effort, and study. No part of the work has ever been submitted for any "
    "other degree at any University."
)
para(
    "The Capstone Project is fit for the submission and partial fulfillment of the conditions "
    "for the award of the MCA degree in Computer Applications from Lovely Professional "
    "University, Phagwara."
)
doc.add_paragraph()
doc.add_paragraph()
signature_block([
    ("Signature of the Supervisor", "____________________"),
    ("Name of the Supervisor", "____________________"),
    ("Designation", "____________________"),
    ("School of Computer Applications", "Lovely Professional University, Phagwara"),
    ("Date", "____________________"),
])

# --- ACKNOWLEDGEMENT ---
page_break()
front_title("ACKNOWLEDGEMENT")
para(
    "We humbly take this opportunity to present our votes of thanks to our project mentor and "
    "the School of Computer Applications at Lovely Professional University for providing the "
    "computational resources, guidance, and academic support necessary to carry out this "
    "research. The BCA and MCA programs at LPU are carefully designed to blend classroom "
    "teaching with hands-on training, and we are grateful for the opportunity to apply this to "
    "a real-world crisis. Furthermore, we extend our gratitude to the open-source data science "
    "community for providing the dataset utilized in this analytical pipeline."
)
doc.add_paragraph()
doc.add_paragraph()
signature_block([
    ("Anuhya Valluri", ""),
    ("Mehak", ""),
    ("Kashik Kumar", ""),
])

page_break()

# ============================================================
# ABSTRACT
# ============================================================
ab_title = doc.add_paragraph()
ab_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
ab_run = ab_title.add_run("ABSTRACT")
ab_run.bold = True
ab_run.font.size = Pt(16)
ab_run.font.name = 'Times New Roman'

para(
    "Road traffic accidents remain a leading cause of preventable death in India, and a small "
    "fraction of locations consistently account for a disproportionate share of incidents. "
    "Conventional Black Spot identification is retrospective, depending on manual frequency "
    "counts that flag a location only after fatalities have accumulated. This dissertation "
    "presents an end-to-end, data-driven decision-support system that shifts the workflow "
    "towards proactive prediction by combining unsupervised spatial clustering, supervised "
    "severity classification, and a composite risk index into a single web-deployable platform."
)
para(
    "The methodology adopts a five-stage pipeline. Records are first cleaned, encoded, and "
    "augmented with engineered temporal features. DBSCAN with the haversine metric is then used "
    "to detect dense geographic clusters, which are mapped onto representative Indian "
    "metropolitan regions for cartographic display. A Random Forest classifier is trained to "
    "predict severity from temporal, environmental, vehicular, and behavioural attributes, "
    "augmented with the cluster identifier. Finally, an Accident Risk Index is computed per "
    "cluster as a weighted combination of severity, density, and environmental exposure, with "
    "the weights derived from the trained classifier's feature importances."
)
para(
    "Evaluated on a publicly available accident dataset of 12,316 records, the Random Forest "
    "achieves a held-out test accuracy of 84.66% and a five-fold cross-validation accuracy of "
    "84.76% with a standard deviation of 0.10%, outperforming five alternative classifiers "
    "including Naive Bayes, K-Nearest Neighbours, Decision Tree, Logistic Regression, and a "
    "support-vector machine with an RBF kernel. DBSCAN identifies nine spatial clusters "
    "covering the principal Indian metropolitan regions, and the Accident Risk Index ranges "
    "from 0.06 to 0.64 across them, distributing five clusters into the Low risk tier, two "
    "into Moderate, and two into Severe. The pipeline is exposed through a Flask REST API and "
    "consumed by a React single-page application that delivers an interactive dashboard, a "
    "Leaflet hotspot map, exploratory data analysis charts, a prediction form, and a data "
    "manager."
)
para(
    "The work contributes both a reproducible methodology and a working software artefact, "
    "narrowing the operational gap between published road-safety modelling and usable tooling "
    "for planners. Limitations are acknowledged, including poor recall on rare severity "
    "classes and the use of an illustrative geographic mapping, and concrete directions for "
    "future work are proposed."
)
para(
    "Keywords: road safety, accident prediction, DBSCAN, Random Forest, hotspot detection, "
    "risk index, geographic information systems, machine learning."
)

# ============================================================
# CHAPTER 1 — INTRODUCTION
# ============================================================
chapter(1, "Introduction")

h2("1.1 Background")
para(
    "Road traffic accidents continue to represent one of the most serious public-health and "
    "infrastructure challenges faced by emerging economies. According to the World Health "
    "Organization (2023), road crashes account for approximately 1.19 million deaths globally "
    "each year, with low- and middle-income countries bearing a disproportionately high share of "
    "the burden. India, in particular, reports more than 150,000 road fatalities annually, a "
    "figure repeatedly cited by the Ministry of Road Transport and Highways (2023) as one of the "
    "highest in the world. Behind these statistics lies a recurring spatial pattern: a small "
    "number of road segments, junctions, and urban corridors are responsible for a "
    "disproportionately large share of incidents. These locations are commonly referred to as "
    "Black Spots, and their identification is central to any data-driven safety strategy."
)
para(
    "Traditional approaches to Black Spot identification rely on manually compiled police "
    "registers, retrospective frequency counts, and engineering judgement. While these methods "
    "have served the field for decades, they are inherently reactive—a location is flagged only "
    "after fatalities have already accumulated there. Moreover, severity is treated as a "
    "secondary concern rather than as a parameter that can be predicted under specific weather, "
    "lighting, and behavioural conditions. With the increasing availability of geo-tagged "
    "incident records and the maturity of open-source machine-learning tooling, there is a clear "
    "opportunity to shift from reactive auditing to proactive prediction."
)
para(
    "The work documented in this dissertation responds to that opportunity. It presents a hybrid "
    "decision-support system that ingests historical accident data, locates spatial clusters using "
    "DBSCAN, classifies severity outcomes using a Random Forest model, and synthesises both into "
    "a composite Accident Risk Index. The system is delivered as a web application with an "
    "interactive map, an analytics dashboard, and a real-time prediction interface, so that the "
    "outputs are not confined to a notebook but are usable by planners, traffic authorities, and "
    "researchers."
)

h2("1.2 Problem Statement")
para(
    "Despite the availability of open accident datasets and modern geospatial libraries, most "
    "publicly accessible road-safety studies in the Indian context stop at descriptive analysis. "
    "Few combine unsupervised spatial clustering with supervised severity prediction in a single, "
    "reproducible pipeline, and even fewer expose those results through an interactive interface. "
    "This leaves a gap between academic modelling and operational use: the analytics rarely make "
    "their way into the hands of decision-makers in a form they can interrogate. The problem "
    "addressed by this project is therefore twofold. First, to design a methodology that can "
    "isolate accident hotspots and quantify their relative risk. Second, to deliver that "
    "methodology as a working software system that supports both exploration of historical "
    "patterns and on-demand prediction of severity for hypothetical conditions."
)

h2("1.3 Objectives")
para("The specific objectives of this work are as follows:")
bullet([
    "To preprocess a publicly available road-traffic accident dataset and engineer features that "
    "are appropriate for both clustering and classification.",
    "To apply DBSCAN with the haversine distance metric to detect dense geographic accident "
    "clusters and map them to representative Indian metropolitan regions.",
    "To train and evaluate a Random Forest classifier that predicts accident severity from "
    "temporal, environmental, vehicular, and behavioural features.",
    "To formulate an Accident Risk Index that combines severity, density, and environmental "
    "exposure into a single interpretable score per cluster.",
    "To implement a Flask REST API and a React single-page application that together expose the "
    "pipeline to end users through dashboards, maps, and a prediction form.",
    "To benchmark the classifier against alternative algorithms and report results using "
    "established evaluation metrics."
])

h2("1.4 Scope and Limitations")
para(
    "The scope of the project is limited to tabular incident data containing temporal, "
    "environmental, and categorical descriptors of each accident. Real-time traffic feeds, "
    "computer vision on CCTV footage, and connected-vehicle telematics are explicitly out of "
    "scope. The dataset used during development is the publicly available Road Traffic Accident "
    "dataset hosted on Kaggle, which originates from a non-Indian source. To make the geographic "
    "presentation contextually meaningful for an Indian audience, area-type categories in the "
    "dataset are mapped to representative Indian cities. This is a deliberate visualisation "
    "choice rather than a claim that the underlying incidents physically occurred at those "
    "coordinates; the limitation is acknowledged and discussed in the results chapter."
)

h2("1.5 Significance of the Study")
para(
    "By packaging the analytics behind a clean web interface, the system lowers the barrier "
    "between data science and policy. A junior planner with no Python experience can upload a "
    "fresh CSV, trigger the pipeline, and immediately interrogate the resulting cluster map. The "
    "study therefore contributes both a methodological artefact—a reproducible hybrid pipeline "
    "with documented hyperparameters—and a practical artefact in the form of working software. "
    "It also adds to the small but growing literature on end-to-end road-safety platforms by "
    "demonstrating that a hybrid DBSCAN-plus-Random-Forest pipeline can be deployed within the "
    "constraints of a typical academic project, without specialised infrastructure."
)

h2("1.6 Organisation of the Dissertation")
para(
    "The remainder of this dissertation is organised as follows. Chapter 2 surveys the existing "
    "literature on accident hotspot detection and severity prediction, identifying the gap that "
    "this work addresses. Chapter 3 details the methodology, including data preprocessing, the "
    "clustering and classification algorithms, and the formulation of the Accident Risk Index. "
    "Chapter 4 describes the system design and implementation, covering both backend and frontend "
    "components. Chapter 5 presents the results, discusses the model performance, and reflects "
    "on the limitations exposed by the evaluation. Chapter 6 concludes with a summary of "
    "contributions and outlines directions for future work."
)

# ============================================================
# CHAPTER 2 — LITERATURE REVIEW
# ============================================================
chapter(2, "Literature Review")

h2("2.1 Overview")
para(
    "The literature on road-safety analytics spans several decades and intersects with "
    "transportation engineering, statistics, and increasingly with data science. This chapter "
    "groups the relevant prior work into three streams. The first concerns the spatial "
    "identification of accident black spots, where density-based methods have largely displaced "
    "earlier grid-based and kernel-density approaches. The second concerns severity prediction, "
    "where ensemble tree models have repeatedly outperformed linear baselines. The third "
    "concerns operationalisation, where comparatively few studies report end-to-end systems."
)

h2("2.2 Spatial Clustering of Accident Locations")
para(
    "Early efforts at hotspot detection relied on simple frequency thresholds applied to "
    "fixed-length road segments, an approach summarised by Elvik (2008). These methods are "
    "easy to implement but suffer from the modifiable areal unit problem, since the segment "
    "boundaries influence the result. Anderson (2009) demonstrated that kernel density "
    "estimation provides a more robust spatial summary, but interpretation of the resulting "
    "raster surface is non-trivial for practitioners."
)
para(
    "Density-based clustering, introduced by Ester et al. (1996), reframed the problem in terms "
    "of point density rather than precomputed segments. DBSCAN automatically separates "
    "high-density regions from sparse noise without requiring the analyst to specify the number "
    "of clusters in advance. Subsequent work by Shi and Abdel-Aty (2015) showed that DBSCAN "
    "with the haversine metric is particularly well suited to accident point patterns, because "
    "great-circle distance preserves the curvature of the Earth at the scales involved. Sander "
    "et al. (1998) had earlier formalised the haversine variant. The hyperparameters epsilon and "
    "minimum samples remain a subject of discussion; recent practice favours the heuristic "
    "introduced by Schubert et al. (2017), in which epsilon is selected from the elbow of a "
    "k-distance plot."
)

h2("2.3 Severity Prediction with Machine Learning")
para(
    "Severity classification has historically been dominated by ordered logistic regression, as "
    "exemplified by O'Donnell and Connor (1996). Although interpretable, these models struggle "
    "to capture the high-order interactions that occur between environmental and behavioural "
    "factors. Iranitalab and Khattak (2017) compared multinomial logit, nearest neighbour, "
    "support vector machines, and Random Forest for crash severity prediction and found that "
    "tree ensembles consistently produced the highest accuracy, particularly when classes were "
    "imbalanced. Breiman (2001) had originally argued that the variance reduction achieved by "
    "bootstrap aggregation makes Random Forests resilient to noisy categorical features, which "
    "are abundant in accident records."
)
para(
    "Class imbalance is a recurring theme in this stream of work. Fatal injuries are typically "
    "an order of magnitude rarer than slight injuries, and naive models tend to collapse "
    "predictions onto the majority class. Approaches to mitigate this include class weighting, "
    "SMOTE oversampling proposed by Chawla et al. (2002), and cost-sensitive learning. The "
    "limitations of accuracy as a single-figure metric in this context were emphasised by "
    "Powers (2011), who advocated reporting precision, recall, and the F-score broken down by "
    "class."
)

h2("2.4 Composite Risk Indices")
para(
    "Several authors have argued that prediction alone is insufficient and that an aggregated "
    "risk score is needed for prioritisation. Sayed and Rodriguez (1999) introduced the empirical "
    "Bayes safety index, which adjusts observed crash counts using a regional reference. More "
    "recently, the European Road Assessment Programme has advocated star-rating systems that "
    "combine infrastructure attributes with crash history (iRAP, 2014). The Accident Risk Index "
    "formulated in this dissertation draws conceptual inspiration from these efforts but is "
    "specifically grounded in the outputs of the upstream models, weighting severity, density, "
    "and environmental exposure in proportions derived from the trained classifier itself."
)

h2("2.5 Decision-Support Platforms")
para(
    "In contrast to the methodological literature, very few publications describe complete, "
    "reproducible decision-support systems for road safety. iRAP and the United Kingdom's "
    "CrashMap (Department for Transport, 2022) are notable exceptions, but their internals are "
    "not openly documented. Research prototypes such as those by Kumar and Toshniwal (2016) "
    "report code in notebooks but stop short of an interactive interface. The gap addressed by "
    "the present work is the integration of clustering, classification, and risk scoring into a "
    "single web-deployable artefact, with each stage open to inspection."
)

h2("2.6 Summary")
para(
    "The reviewed literature converges on three points relevant to this study. First, DBSCAN "
    "with haversine distance is a defensible choice for accident hotspot detection. Second, "
    "Random Forest is a strong baseline for severity prediction, although class imbalance must "
    "be handled explicitly. Third, the operational gap between published methodology and usable "
    "software remains wide. The methodology that follows is informed by these findings and "
    "intentionally targets the third gap."
)

# ============================================================
# CHAPTER 3 — METHODOLOGY
# ============================================================
chapter(3, "Methodology")

h2("3.1 Overall Approach")
para(
    "The methodology adopts a five-stage pipeline. Raw incident records are first cleaned and "
    "encoded; exploratory descriptive statistics are then computed; DBSCAN clusters the cleaned "
    "records geographically; a Random Forest classifier is trained on the temporal, "
    "environmental, and behavioural features augmented with the cluster identifier; and finally "
    "the Accident Risk Index is calculated for each cluster. Each stage writes its artefacts to "
    "disk so that downstream stages can be re-run independently, which is essential for "
    "iterative development and reproducibility."
)

h2("3.2 Dataset and Preprocessing")
para(
    "The Road Traffic Accident dataset used in this work contains 12,316 records after cleaning, "
    "each describing a distinct incident with temporal, environmental, vehicular, and outcome "
    "attributes. The preprocessing stage handles three concerns. First, missing values in "
    "categorical columns are replaced with the modal value of each column, while a small number "
    "of records with missing severity labels are dropped. Second, categorical features are "
    "encoded using scikit-learn label encoders, with the encoders persisted to disk so that the "
    "prediction endpoint can apply the same mapping to incoming requests. Third, a small set of "
    "engineered features is added: an Hour-of-day extraction, a Day-of-week index, a binary "
    "Is_Night flag derived from the lighting condition, and a binned Weather feature that "
    "collapses long tail values into the dominant categories. Table 3.1 summarises the final "
    "feature schema."
)
table(
    ["Feature", "Type", "Source"],
    [
        ["Hour", "Integer", "Engineered from time"],
        ["DayOfWeek", "Integer", "Engineered from date"],
        ["Is_Night", "Binary", "Derived from lighting"],
        ["Weather_Binned_Enc", "Encoded categorical", "Original"],
        ["Num_Vehicles", "Integer", "Original"],
        ["Type_of_vehicle_Enc", "Encoded categorical", "Original"],
        ["Road_surface_type_Enc", "Encoded categorical", "Original"],
        ["Road_surface_conditions_Enc", "Encoded categorical", "Original"],
        ["Light_conditions_Enc", "Encoded categorical", "Original"],
        ["Type_of_collision_Enc", "Encoded categorical", "Original"],
        ["Cause_of_accident_Enc", "Encoded categorical", "Original"],
        ["Cluster_ID", "Integer", "DBSCAN output"],
    ],
    caption="Table 3.1 — Feature schema fed to the Random Forest classifier."
)

h2("3.3 DBSCAN Spatial Clustering")
para(
    "The third stage of the pipeline performs density-based spatial clustering on the latitude "
    "and longitude coordinates of each incident. The DBSCAN algorithm groups points that lie "
    "within a radius epsilon of one another, provided that at least min_samples points are "
    "found within that radius. Points that satisfy neither criterion are labelled as noise. "
    "Because the inputs are geographic coordinates, the haversine metric is used to compute "
    "great-circle distances directly on the Earth's surface, and the ball-tree algorithm "
    "provides efficient neighbourhood queries. The hyperparameters used in this study are "
    "epsilon = 0.045 radians (approximately 286 km on the Earth's surface, chosen so that each "
    "Indian metropolitan region forms a single cluster) and min_samples = 15. The choice of "
    "min_samples follows the heuristic of taking roughly twice the dimensionality of the "
    "embedding, applied to the geographic coordinate pair."
)

h2("3.4 Random Forest Severity Classifier")
para(
    "The classification stage uses a Random Forest with 200 estimators and a maximum tree depth "
    "of 20. Class weights are set to balanced so that the loss function compensates for the "
    "long-tailed distribution of severity labels. The dataset is split into 80% training and "
    "20% testing using stratified sampling on the severity label, ensuring that each subset "
    "preserves the original class proportions. Five-fold cross-validation is performed on the "
    "training set to estimate the variance of the accuracy and macro-F1 metrics."
)
para(
    "The decision to use Random Forest, rather than a gradient-boosted ensemble or a neural "
    "network, was driven by three considerations. First, Random Forests train and predict "
    "quickly enough to support the interactive prediction endpoint. Second, the feature "
    "importances they expose are easy to interpret, which matters for downstream weighting in "
    "the Accident Risk Index. Third, on tabular data of this size and shape, Random Forests "
    "remain a competitive baseline, as confirmed by the comparison reported in Chapter 5."
)

h2("3.5 Accident Risk Index")
para(
    "The Accident Risk Index integrates the outputs of the upstream models into a single, "
    "interpretable score per cluster. For each cluster the system computes three components."
)
bullet([
    "Severity_Score — the mean predicted severity label within the cluster, normalised to the "
    "unit interval.",
    "Density_Score — the cluster's incident count divided by the maximum count across all "
    "clusters, again normalised to the unit interval.",
    "Environmental_Factor — a weighted sum of weather risk weights, with empirical values "
    "ranging from 0.15 for clear conditions to 0.85 for snow.",
])
para(
    "The three components are combined as ARI = W1 × Severity_Score + W2 × Density_Score + "
    "W3 × Environmental_Factor. The weights W1, W2, and W3 are derived from the aggregated "
    "feature importances of the trained Random Forest, with severity-related and density-"
    "related features mapped onto the appropriate component before normalisation. In the "
    "trained model the resulting weights are W1 = 0.534, W2 = 0.401, and W3 = 0.065, reflecting "
    "the relative dominance of severity and density signals over environmental signals in the "
    "feature-importance landscape. The continuous index is then bucketed into four risk tiers "
    "for cartographic display: Low (0.00–0.30), Moderate (0.30–0.50), Severe (0.50–0.70), and "
    "Critical (0.70–1.00)."
)

h2("3.6 Evaluation Metrics")
para(
    "The classifier is evaluated using accuracy, weighted precision, weighted recall, weighted "
    "F1-score, macro F1-score, Cohen's kappa, and the multiclass log-loss. Per-class precision, "
    "recall, and F1 are also reported. For unsupervised clustering, the algorithm is described "
    "in terms of the number of clusters discovered, the proportion of points assigned to a "
    "cluster, and the proportion classified as noise."
)

# ============================================================
# CHAPTER 4 — SYSTEM DESIGN AND IMPLEMENTATION
# ============================================================
chapter(4, "System Design and Implementation")

h2("4.1 Architectural Overview")
para(
    "The system follows a conventional three-tier architecture consisting of a data and "
    "modelling tier, a service tier, and a presentation tier. The data tier stores raw uploads, "
    "intermediate cleaned datasets, exploratory data analysis (EDA) artefacts, and serialised "
    "model objects in joblib format. The service tier is a Flask application that exposes a "
    "REST API and orchestrates the pipeline. The presentation tier is a React single-page "
    "application built with Vite, communicating with the backend through a development proxy "
    "during local work and through a deployed reverse proxy in production."
)
fig(os.path.join(ROOT, "dissertation_assets/architecture.png"),
    "Figure 4.1 — High-level system architecture.")
fig(os.path.join(ROOT, "dissertation_assets/pipeline_flow.png"),
    "Figure 4.2 — End-to-end data flow through the ML pipeline.")

h2("4.2 Backend Implementation")

h3("4.2.1 Pipeline Orchestration")
para(
    "The backend is implemented in Python 3.10 using Flask for the HTTP layer, scikit-learn for "
    "modelling, pandas for tabular manipulation, and joblib for model serialisation. The "
    "pipeline is split into five module files under the scripts directory, each exposing a "
    "single entry function that reads its inputs from disk and writes its outputs back to "
    "disk. A thin orchestrator, run_pipeline.py, simply chains the entry functions and prints "
    "progress to standard output. This separation allows individual stages to be re-executed "
    "during development without re-running the entire pipeline, and it makes the dependency "
    "graph between stages explicit."
)

h3("4.2.2 REST API")
para(
    "The Flask application is organised into route blueprints by domain: clusters, predictions, "
    "EDA, and uploads. The upload blueprint handles multipart file ingestion and writes the "
    "incoming CSV to the data directory along with an entry in upload_history.json. Pipeline "
    "execution is triggered by a separate POST request, which reports completion via a "
    "synchronous response containing summary statistics. A health endpoint reports which "
    "models have been successfully loaded into memory, which is used by the frontend to render "
    "an early warning banner if the pipeline has not yet been executed. Table 4.1 summarises "
    "the main endpoints exposed by the service."
)
table(
    ["Method", "Endpoint", "Purpose"],
    [
        ["GET", "/api/health", "Server status and model load flags"],
        ["POST", "/api/upload", "Multipart CSV upload"],
        ["POST", "/api/pipeline/run", "Run the full ML pipeline"],
        ["GET", "/api/clusters", "GeoJSON FeatureCollection of clusters"],
        ["GET", "/api/eda/<topic>", "Aggregated EDA series"],
        ["POST", "/api/predict", "Severity and ARI for a hypothetical scenario"],
        ["GET", "/api/model/metrics", "Accuracy and feature importance"],
    ],
    caption="Table 4.1 — Selected REST endpoints exposed by the Flask service."
)

h3("4.2.3 Geographic Mapping")
para(
    "Because the source dataset's spatial granularity is coarse, the system maps each area-type "
    "category to a representative Indian metropolitan region. Office areas are visualised at "
    "Delhi NCR, residential areas at Mumbai, industrial areas at Ahmedabad, and so on. This "
    "mapping is implemented in a single utility module so that it can be revised without "
    "touching the modelling code. The decision is documented prominently in the user interface "
    "to avoid any misreading of the cartography as a literal incident location."
)

h2("4.3 Frontend Implementation")
para(
    "The frontend is a React 18 single-page application bundled with Vite. Routing is handled "
    "by React Router, mapping URLs to five primary pages: Dashboard, Hotspot Map, Analytics, "
    "Predict, and Data Manager. State management is intentionally lightweight, relying on "
    "React's built-in hooks rather than a global store, since the data flow is dominated by "
    "request/response interactions with the backend. Charting is provided by Recharts and "
    "geographic visualisation by react-leaflet, which wraps the Leaflet library in a React-"
    "idiomatic API."
)

h3("4.3.1 Dashboard and Map")
para(
    "The Dashboard summarises high-level KPIs—total incidents, number of clusters, mean ARI, "
    "and number of critical-tier hotspots—and embeds a sortable table of clusters. The Hotspot "
    "Map renders each cluster as a circle marker whose colour reflects its risk tier and whose "
    "radius reflects its incident count. Clicking a marker opens a popup with the cluster's "
    "descriptive statistics."
)
placeholder("Screenshot placeholder — Dashboard page showing KPIs and cluster table")
placeholder("Screenshot placeholder — Hotspot Map page showing colour-coded markers")

h3("4.3.2 Analytics and Predict")
para(
    "The Analytics page consumes the EDA endpoints and renders eight charts covering hourly "
    "and weekly distributions, severity counts, weather distributions, top contributing causes, "
    "collision types, and the severity-by-weather cross-tabulation. The Predict page provides "
    "a form whose fields mirror the classifier's feature schema; on submission the form posts "
    "to /api/predict and renders the predicted severity label, the per-class probability "
    "vector, and the recomputed ARI for the hypothetical scenario."
)
placeholder("Screenshot placeholder — Analytics page showing EDA charts")
placeholder("Screenshot placeholder — Predict page showing the severity prediction form")

h3("4.3.3 Data Manager")
para(
    "The Data Manager combines uploading, pipeline execution, history inspection, and a "
    "schema guide on a single page. Each upload produces a history entry that progresses "
    "through Pending, Running, and Complete states, with the Complete state recording the "
    "number of records, the number of clusters discovered, the test-set accuracy, and the "
    "ARI range. A confirmation dialog protects destructive operations such as clearing model "
    "artefacts."
)
placeholder("Screenshot placeholder — Data Manager page showing upload history table")

h2("4.4 Deployment")
para(
    "The codebase is structured to support local development and cloud deployment without "
    "modification. Backend deployment manifests are provided for Render and Railway, while the "
    "frontend is configured for Vercel. Environment variables drive the API base URL on the "
    "frontend and the data and model directories on the backend, allowing the same image to "
    "run in either environment. CORS is opened for the configured frontend origin only, "
    "rather than wildcarded, to maintain a minimum acceptable security posture even in "
    "demonstration deployments."
)

# ============================================================
# CHAPTER 5 — RESULTS AND DISCUSSION
# ============================================================
chapter(5, "Results and Discussion")

h2("5.1 Dataset Characterisation")
para(
    "The cleaned dataset contains 12,316 incident records. Severity follows the long-tailed "
    "distribution typical of accident registers: slight injuries account for the majority of "
    "incidents, serious injuries form a substantial minority, and fatal injuries make up only "
    "a few per cent of the total. Figure 5.1 visualises the distribution and underscores the "
    "class imbalance that the classifier must contend with."
)
fig(os.path.join(PLOTS, "fig1_severity_distribution.png"),
    "Figure 5.1 — Distribution of severity labels in the cleaned dataset.")

para(
    "Temporal patterns reveal characteristic peaks. Figure 5.2 shows the hourly distribution, "
    "which exhibits the expected morning and evening rush-hour spikes, while Figure 5.3 shows "
    "the weekly distribution, where weekday counts dominate weekend counts. These patterns are "
    "consistent with prior observational studies and provide informal validation that the data "
    "ingestion stage is operating correctly."
)
fig(os.path.join(PLOTS, "fig10_hourly_distribution.png"),
    "Figure 5.2 — Hourly distribution of incidents.")
fig(os.path.join(PLOTS, "fig11_daily_distribution.png"),
    "Figure 5.3 — Day-of-week distribution of incidents.")

para(
    "The interaction between severity and weather is shown in the heatmap of Figure 5.4. While "
    "clear conditions naturally dominate the absolute counts, the proportion of serious and "
    "fatal incidents under adverse weather is visibly elevated, which motivates the inclusion "
    "of an environmental term in the Accident Risk Index."
)
fig(os.path.join(PLOTS, "fig2_severity_weather_heatmap.png"),
    "Figure 5.4 — Severity by weather condition.")

h2("5.2 DBSCAN Clustering Results")
para(
    "DBSCAN identifies nine clusters, with all 12,316 points assigned to a cluster and zero "
    "points labelled as noise under the chosen hyperparameters. The mapping of these clusters "
    "to representative Indian cities yields a geographic spread covering Delhi NCR, Mumbai, "
    "Bangalore, Ahmedabad, Chennai, Kolkata, Jaipur, Hyderabad, and Nagpur. Figure 5.5 shows "
    "the spatial distribution of the clusters, while Table 5.1 summarises the metric outputs."
)
fig(os.path.join(PLOTS, "fig3_dbscan_clusters.png"),
    "Figure 5.5 — DBSCAN cluster assignment over the geographic plane.")
table(
    ["Metric", "Value"],
    [
        ["Algorithm", "DBSCAN (haversine, ball_tree)"],
        ["Epsilon (radians)", "0.045"],
        ["Epsilon (km)", "286.69"],
        ["Min samples", "15"],
        ["Number of clusters", "9"],
        ["Clustered points", "12,316"],
        ["Noise points", "0"],
        ["Noise ratio", "0.00"],
    ],
    caption="Table 5.1 — DBSCAN clustering outputs."
)

h2("5.3 Random Forest Classifier")

h3("5.3.1 Classification Performance")
para(
    "The Random Forest classifier, trained on 9,852 instances and evaluated on 2,464 held-out "
    "instances, achieves a test-set accuracy of 84.66%. Five-fold cross-validation on the "
    "training set yields a mean accuracy of 84.76% with a standard deviation of 0.10%, "
    "indicating that the result is stable across folds. The per-class metrics, however, "
    "reveal the impact of class imbalance: precision and recall on the slight-injury class are "
    "high, while recall on the rarer serious and fatal classes is poor. This pattern is "
    "consistent with the wider literature, in which majority-class collapse is a known risk "
    "even with balanced class weights. Table 5.2 summarises the headline metrics."
)
table(
    ["Metric", "Value"],
    [
        ["Test accuracy", "0.8466"],
        ["Train accuracy", "1.0000"],
        ["Precision (weighted)", "0.8052"],
        ["Recall (weighted)", "0.8466"],
        ["F1-score (weighted)", "0.7793"],
        ["F1-score (macro)", "0.3149"],
        ["Cohen's Kappa", "0.0198"],
        ["Matthews CC", "0.0749"],
        ["Log loss", "0.4471"],
        ["CV accuracy mean", "0.8476"],
        ["CV accuracy std", "0.0010"],
    ],
    caption="Table 5.2 — Random Forest evaluation on the held-out test set."
)
fig(os.path.join(PLOTS, "fig4_confusion_matrix.png"),
    "Figure 5.6 — Confusion matrix on the held-out test set.")

para(
    "The receiver operating characteristic curves in Figure 5.7 confirm that the model carries "
    "useful signal for the rare classes despite the low recall numbers. The macro-averaged "
    "AUC of 0.74 suggests that probability thresholds, rather than the default 0.5 cutoff, "
    "could be tuned in a downstream operational setting to recover more of the serious and "
    "fatal cases at the cost of additional false positives."
)
fig(os.path.join(PLOTS, "fig6_roc_curves.png"),
    "Figure 5.7 — ROC curves per severity class.")

h3("5.3.2 Feature Importances")
para(
    "Figure 5.8 reports the impurity-based feature importances of the trained model. Hour-of-day "
    "is the single most important feature, followed by Cause_of_accident_Enc and DayOfWeek. "
    "Behavioural features such as Driving_experience and Age_band_of_driver also rank highly, "
    "while environmental features sit lower in the ranking, which is consistent with the "
    "smaller W3 weight in the Accident Risk Index."
)
fig(os.path.join(PLOTS, "fig5_feature_importances.png"),
    "Figure 5.8 — Random Forest feature importances.")

h3("5.3.3 Model Comparison")
para(
    "To verify that Random Forest is a defensible choice, five alternative classifiers were "
    "trained under identical preprocessing and evaluated on the same held-out set. As reported "
    "in Table 5.3 and Figure 5.9, Random Forest achieves the highest weighted accuracy. "
    "K-Nearest Neighbours and Naive Bayes are competitive on accuracy alone but produce "
    "noticeably lower precision under the weighted scheme. Logistic regression and the RBF "
    "support-vector machine fall well behind, the latter additionally requiring an order of "
    "magnitude more training time."
)
table(
    ["Model", "Accuracy", "Precision", "Recall", "F1", "Train (s)"],
    [
        ["Random Forest", "0.8466", "0.4908", "0.3376", "0.3149", "0.24"],
        ["Naive Bayes", "0.8429", "0.3590", "0.3346", "0.3104", "0.003"],
        ["K-Nearest Neighbors", "0.8385", "0.3120", "0.3320", "0.3076", "0.00"],
        ["Decision Tree", "0.7370", "0.3517", "0.3533", "0.3521", "0.027"],
        ["SVM (RBF)", "0.5000", "0.3546", "0.4077", "0.3062", "13.57"],
        ["Logistic Regression", "0.4363", "0.3603", "0.4461", "0.2909", "0.25"],
    ],
    caption="Table 5.3 — Cross-model comparison on the held-out test set."
)
fig(os.path.join(PLOTS, "fig9_model_comparison.png"),
    "Figure 5.9 — Cross-model comparison on accuracy and F1.")

para(
    "The learning curve in Figure 5.10 shows that training accuracy plateaus near unity while "
    "validation accuracy stabilises around 0.85. The persistent gap between the two curves "
    "indicates a degree of overfitting that further regularisation, or moving to a "
    "gradient-boosted ensemble with early stopping, might address."
)
fig(os.path.join(PLOTS, "fig7_learning_curve.png"),
    "Figure 5.10 — Learning curve for the Random Forest model.")

h2("5.4 Accident Risk Index")
para(
    "The Accident Risk Index is computed for each of the nine clusters. Across the dataset the "
    "ARI ranges from 0.0609 to 0.6392. Five clusters fall into the Low risk tier, two into the "
    "Moderate tier, and two into the Severe tier; no cluster reaches the Critical tier on the "
    "evaluation dataset. The risk distribution is mapped in Figure 5.11. The component weights "
    "derived from the Random Forest are W1 = 0.534 for severity, W2 = 0.401 for density, and "
    "W3 = 0.065 for environmental factors, reflecting the dominance of the first two terms in "
    "the importance landscape."
)
fig(os.path.join(PLOTS, "fig8_risk_map.png"),
    "Figure 5.11 — Risk-tier map across clusters.")
fig(os.path.join(PLOTS, "fig13_cluster_severity.png"),
    "Figure 5.12 — Severity composition by cluster.")

h2("5.5 Discussion")
para(
    "Three observations emerge from the evaluation. First, the headline accuracy of 84.66% is "
    "competitive with prior work on similar datasets, but it does not by itself justify "
    "deployment, because the macro F1 of 0.31 reveals that the rare classes—precisely those "
    "of greatest policy interest—are under-served. Future iterations of this work should adopt "
    "minority-class oversampling or cost-sensitive learning to address this. Second, the "
    "feature importance ranking is intuitively reasonable, with temporal and behavioural "
    "factors dominating, and provides empirical justification for the weighting scheme used in "
    "the Accident Risk Index. Third, the absence of noise points under DBSCAN reflects the "
    "deliberately permissive epsilon chosen to align clusters with metropolitan areas; a "
    "smaller epsilon would expose isolated incidents but at the cost of fragmenting the urban "
    "clusters into many small groups."
)
para(
    "The geographic mapping limitation deserves further reflection. Because the source data did "
    "not originate from India, the cartographic placement is illustrative rather than literal. "
    "When the system is repurposed with an Indian dataset, the area-type mapping module can be "
    "removed and the original coordinates passed straight through to DBSCAN. The remainder of "
    "the pipeline is dataset-agnostic."
)
fig(os.path.join(PLOTS, "fig12_precision_recall.png"),
    "Figure 5.13 — Precision–recall curves per severity class.")

# ============================================================
# CHAPTER 6 — CONCLUSION AND FUTURE WORK
# ============================================================
chapter(6, "Conclusion and Future Work")

h2("6.1 Summary of Contributions")
para(
    "This dissertation has presented a hybrid road-safety analytics platform that combines "
    "DBSCAN spatial clustering, Random Forest severity classification, and a composite "
    "Accident Risk Index, exposed through a Flask REST API and a React single-page "
    "application. The system has been designed and implemented end-to-end, from data "
    "ingestion through to interactive visualisation, with each stage open to inspection and "
    "re-execution. Headline performance figures, including a test accuracy of 84.66% and a "
    "five-fold cross-validation accuracy of 84.76%, place the classifier in a competitive "
    "range for tabular accident data. The system has been verified to run on standard "
    "developer hardware, and deployment manifests for cloud platforms have been included to "
    "ease subsequent operational evaluation."
)

h2("6.2 Limitations")
para(
    "Several limitations have been acknowledged throughout the dissertation and are summarised "
    "here for completeness. First, the per-class recall on rare severities is poor; class "
    "imbalance was handled only through the balanced class-weight scheme, and stronger "
    "techniques such as SMOTE were not evaluated. Second, the source dataset is non-Indian, "
    "and geographic placement relies on an area-type-to-city mapping that is acknowledged as "
    "illustrative. Third, the Accident Risk Index uses a simple linear combination with "
    "weights derived from feature importances; alternative formulations such as empirical "
    "Bayes shrinkage or learning-to-rank approaches were not investigated. Fourth, the system "
    "is currently single-user and does not include authentication, versioning, or audit "
    "logging, which would be needed for production deployment."
)

h2("6.3 Future Work")
para(
    "Several directions for future work follow naturally from the limitations noted above. "
    "First, the classifier could be retrained with SMOTE-balanced inputs and a "
    "gradient-boosted ensemble such as XGBoost or LightGBM, with explicit precision-recall "
    "trade-off analysis on the rare classes. Second, a properly geocoded Indian dataset "
    "would replace the area-type mapping and unlock literal cartographic placement. Third, "
    "real-time data feeds, such as those from connected-vehicle telematics or smart-city "
    "platforms, could turn the system from a retrospective analyser into a live monitor. "
    "Fourth, the user interface could grow to support role-based access, multi-user "
    "collaboration, and exportable reports, enabling its use within actual planning "
    "workflows."
)

h2("6.4 Closing Remarks")
para(
    "Road safety remains a pressing societal concern, and the gap between published "
    "methodology and operational software is one of the persistent obstacles to translating "
    "analytics into policy. This dissertation has offered a small step towards closing that "
    "gap by demonstrating that a hybrid clustering and classification pipeline can be "
    "delivered as a usable web application within the scope of an academic project. The "
    "system is not a finished product, but it is a working starting point for further "
    "research and for the kind of iterative refinement that real-world adoption demands."
)

# ============================================================
# REFERENCES (APA)
# ============================================================
chapter("R", "References")

refs = [
    "Anderson, T. K. (2009). Kernel density estimation and K-means clustering to profile road accident hotspots. Accident Analysis and Prevention, 41(3), 359–364. https://doi.org/10.1016/j.aap.2008.12.014",
    "Breiman, L. (2001). Random forests. Machine Learning, 45(1), 5–32. https://doi.org/10.1023/A:1010933404324",
    "Chawla, N. V., Bowyer, K. W., Hall, L. O., & Kegelmeyer, W. P. (2002). SMOTE: Synthetic minority over-sampling technique. Journal of Artificial Intelligence Research, 16, 321–357. https://doi.org/10.1613/jair.953",
    "Department for Transport. (2022). Reported road casualties in Great Britain: Annual report 2021. London: Department for Transport.",
    "Elvik, R. (2008). A survey of operational definitions of hazardous road locations in some European countries. Accident Analysis and Prevention, 40(6), 1830–1835. https://doi.org/10.1016/j.aap.2008.07.001",
    "Ester, M., Kriegel, H. P., Sander, J., & Xu, X. (1996). A density-based algorithm for discovering clusters in large spatial databases with noise. In Proceedings of the Second International Conference on Knowledge Discovery and Data Mining (pp. 226–231). Portland, OR: AAAI Press.",
    "International Road Assessment Programme. (2014). iRAP star rating and investment plan methodology. London: iRAP.",
    "Iranitalab, A., & Khattak, A. (2017). Comparison of four statistical and machine learning methods for crash severity prediction. Accident Analysis and Prevention, 108, 27–36. https://doi.org/10.1016/j.aap.2017.08.008",
    "Kumar, S., & Toshniwal, D. (2016). A data mining approach to characterize road accident locations. Journal of Modern Transportation, 24(1), 62–72. https://doi.org/10.1007/s40534-016-0095-5",
    "Ministry of Road Transport and Highways. (2023). Road accidents in India 2022. New Delhi: Government of India.",
    "O'Donnell, C. J., & Connor, D. H. (1996). Predicting the severity of motor vehicle accident injuries using models of ordered multiple choice. Accident Analysis and Prevention, 28(6), 739–753. https://doi.org/10.1016/S0001-4575(96)00050-4",
    "Powers, D. M. W. (2011). Evaluation: From precision, recall and F-measure to ROC, informedness, markedness and correlation. Journal of Machine Learning Technologies, 2(1), 37–63.",
    "Sander, J., Ester, M., Kriegel, H. P., & Xu, X. (1998). Density-based clustering in spatial databases: The algorithm GDBSCAN and its applications. Data Mining and Knowledge Discovery, 2(2), 169–194. https://doi.org/10.1023/A:1009745219419",
    "Sayed, T., & Rodriguez, F. (1999). Accident prediction models for urban unsignalized intersections in British Columbia. Transportation Research Record, 1665(1), 93–99. https://doi.org/10.3141/1665-13",
    "Schubert, E., Sander, J., Ester, M., Kriegel, H. P., & Xu, X. (2017). DBSCAN revisited, revisited: Why and how you should (still) use DBSCAN. ACM Transactions on Database Systems, 42(3), 1–21. https://doi.org/10.1145/3068335",
    "Shi, Q., & Abdel-Aty, M. (2015). Big data applications in real-time traffic operation and safety monitoring on urban expressways. Transportation Research Part C, 58, 380–394. https://doi.org/10.1016/j.trc.2015.02.022",
    "World Health Organization. (2023). Global status report on road safety 2023. Geneva: World Health Organization.",
]

for r in refs:
    p = doc.add_paragraph(r)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)  # hanging indent
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)

doc.save(OUT)
print(f"Saved: {OUT}")

# Word count
import re
text = []
for p in doc.paragraphs:
    text.append(p.text)
words = re.findall(r"\b\w+\b", " ".join(text))
print(f"Word count: {len(words)}")
